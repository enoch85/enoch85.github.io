from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re
import shutil
from typing import Dict, List
try:
    from deep_translator import GoogleTranslator
    HAS_TRANSLATOR = True
except ImportError:
    HAS_TRANSLATOR = False
    print("⚠ deep_translator not installed. Install with: pip install deep-translator")
    print("  Will use fallback translations.")

def translate_text(text: str, source_lang: str = 'sv', target_lang: str = 'en') -> str:
    """
    Translate text using Google Translate via deep_translator.
    Falls back to original text if translation fails.
    """
    if not HAS_TRANSLATOR:
        return text
    
    try:
        translator = GoogleTranslator(source=source_lang, target=target_lang)
        # Split long texts into chunks (Google Translate has a 5000 char limit)
        if len(text) > 4500:
            # Split by sentences
            sentences = re.split(r'([.!?]\s+)', text)
            translated_parts = []
            current_chunk = ""
            
            for i, part in enumerate(sentences):
                if len(current_chunk + part) < 4500:
                    current_chunk += part
                else:
                    if current_chunk:
                        translated_parts.append(translator.translate(current_chunk))
                    current_chunk = part
            
            if current_chunk:
                translated_parts.append(translator.translate(current_chunk))
            
            return ''.join(translated_parts)
        else:
            return translator.translate(text)
    except Exception as e:
        print(f"  ⚠ Translation failed: {e}")
        return text

def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to a paragraph.
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = OxmlElement('w:r')
    
    # Set the run's style to hyperlink style
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    
    # Create text node
    new_run.text = text
    hyperlink.append(new_run)

    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)

    return hyperlink

def extract_tsx_data(content: str, var_name: str) -> List[Dict]:
    """
    Extract data from TSX file by finding the array assigned to a variable.
    This is a simple parser that handles the structure in our component files.
    """
    # Find the variable declaration
    pattern = rf'const {var_name}:\s*\w+\[\]\s*=\s*\[(.*?)\];'
    match = re.search(pattern, content, re.DOTALL)
    
    if not match:
        return []
    
    array_content = match.group(1)
    items = []
    
    # Split by objects (find { ... }, pairs)
    brace_count = 0
    current_obj = ""
    in_string = False
    escape_next = False
    
    for char in array_content:
        if escape_next:
            current_obj += char
            escape_next = False
            continue
            
        if char == '\\':
            escape_next = True
            current_obj += char
            continue
            
        if char in ['"', "'", '`'] and not escape_next:
            in_string = not in_string
            
        if not in_string:
            if char == '{':
                brace_count += 1
            elif char == '}':
                brace_count -= 1
                
        current_obj += char
        
        if brace_count == 0 and current_obj.strip().endswith('}'):
            # Parse this object
            obj_str = current_obj.strip().rstrip(',').strip()
            if obj_str:
                items.append(parse_tsx_object(obj_str))
            current_obj = ""
    
    return items

def parse_tsx_object(obj_str: str) -> Dict:
    """
    Parse a simple TypeScript object into a Python dict.
    Handles strings, arrays, and nested structures.
    """
    result = {}
    
    # Remove outer braces
    obj_str = obj_str.strip()
    if obj_str.startswith('{') and obj_str.endswith('}'):
        obj_str = obj_str[1:-1]
    
    # Simple regex patterns for key-value pairs
    # Match: key: 'value' or key: "value" or key: `value`
    string_pattern = r"(\w+):\s*['\"`](.*?)['\"`](?=\s*,|\s*}|\s*$)"
    # Match: key: [...] for arrays
    array_pattern = r"(\w+):\s*\[(.*?)\](?=\s*,|\s*}|\s*$)"
    # Match: key: true/false
    bool_pattern = r"(\w+):\s*(true|false)(?=\s*,|\s*}|\s*$)"
    
    # Extract string values
    for match in re.finditer(string_pattern, obj_str, re.DOTALL):
        key = match.group(1)
        value = match.group(2).replace('\\n', '\n').replace("\\'", "'").replace('\\"', '"')
        result[key] = value
    
    # Extract array values
    for match in re.finditer(array_pattern, obj_str, re.DOTALL):
        key = match.group(1)
        array_content = match.group(2)
        # Split by comma (simple approach)
        items = []
        for item in re.findall(r"['\"`](.*?)['\"`]", array_content):
            items.append(item)
        result[key] = items
    
    # Extract boolean values
    for match in re.finditer(bool_pattern, obj_str):
        key = match.group(1)
        value = match.group(2) == 'true'
        result[key] = value
    
    return result

def extract_content_from_repo() -> Dict:
    """
    Extract content directly from the repository's TSX component files.
    This is more reliable than web scraping.
    """
    print("\n📂 Reading content from repository files...")
    
    try:
        # Get the script's directory and navigate to workspace root
        script_dir = os.path.dirname(os.path.abspath(__file__))
        repo_root = os.path.dirname(script_dir)
        
        data = {
            'profile_text_sv': '',
            'profile_text_en': '',
            'profile_title_sv': '',
            'profile_title_en': '',
            'experiences': [],
            'skills': [],
            'projects': [],
            'logos': {}
        }
        
        # Read Hero.tsx for profile text and title
        hero_path = os.path.join(repo_root, 'components', 'Hero.tsx')
        if os.path.exists(hero_path):
            with open(hero_path, 'r', encoding='utf-8') as f:
                content = f.read()
                
                # Extract the title (subtitle line)
                title_match = re.search(r'<p className="text-xl.*?>\s*(.+?)\s*</p>', content, re.DOTALL)
                if title_match:
                    data['profile_title_sv'] = re.sub(r'\s+', ' ', title_match.group(1)).strip()
                    data['profile_title_en'] = translate_text(data['profile_title_sv'])
                    print(f"  ✓ Found profile title")
                
                # Extract the Swedish profile text - look for the max-w-2xl paragraph
                desc_match = re.search(r'<p className="max-w-2xl[^>]*>\s*(.+?)\s*</p>', content, re.DOTALL)
                if desc_match:
                    # Extract the text and clean up whitespace
                    raw_text = desc_match.group(1)
                    # Remove extra whitespace and newlines but keep sentence structure
                    data['profile_text_sv'] = re.sub(r'\s+', ' ', raw_text).strip()
                    print(f"  ✓ Found profile text (Swedish)")
                    if data['profile_text_sv']:
                        data['profile_text_en'] = translate_text(data['profile_text_sv'])
                        print(f"  ✓ Translated profile text to English")
                    else:
                        print(f"  ⚠ Profile text was empty after extraction")
        
        # Read Experience.tsx for work history
        exp_path = os.path.join(repo_root, 'components', 'Experience.tsx')
        if os.path.exists(exp_path):
            with open(exp_path, 'r', encoding='utf-8') as f:
                content = f.read()
                experiences = extract_tsx_data(content, 'experiences')
                data['experiences'] = experiences
                print(f"  ✓ Extracted {len(experiences)} work experiences")
        
        # Read Skills.tsx for skills
        skills_path = os.path.join(repo_root, 'components', 'Skills.tsx')
        if os.path.exists(skills_path):
            with open(skills_path, 'r', encoding='utf-8') as f:
                content = f.read()
                skills = extract_tsx_data(content, 'skillCategories')
                data['skills'] = skills
                print(f"  ✓ Extracted {len(skills)} skill categories")
        
        # Read Projects.tsx for projects
        projects_path = os.path.join(repo_root, 'components', 'Projects.tsx')
        if os.path.exists(projects_path):
            with open(projects_path, 'r', encoding='utf-8') as f:
                content = f.read()
                projects = extract_tsx_data(content, 'projects')
                data['projects'] = projects
                print(f"  ✓ Extracted {len(projects)} projects")
        
        # Copy logos from public folder to CV directory
        print("\n📥 Copying company logos...")
        public_dir = os.path.join(repo_root, 'public')
        # Use /tmp for dev containers, or the default path if it exists
        default_logos_dir = '/home/daniel/Music/CV/logos'
        if os.path.exists('/home/daniel'):
            logos_dir = default_logos_dir
        else:
            # In dev container or other environment, use a temp directory
            logos_dir = os.path.join(repo_root, 'cv-generator', 'logos')
        os.makedirs(logos_dir, exist_ok=True)
        
        logo_files = {
            'InfraCom': 'infracom-logo.png',
            'Redpill Linpro': 'redpill-linpro-logo.png',
            'T&M Hansson IT': 'hanssonit-logo.png',
            'SenseNode': 'sensenode-logo.png',
            'Vessinge IT': 'vessinge-logo.png',
            'Orestad Linux': 'orestad-logo.png',
            'HSN Konsult': 'hsn-konsult-logo.jpg'
        }
        
        for company_key, logo_file in logo_files.items():
            src_path = os.path.join(public_dir, logo_file)
            dst_path = os.path.join(logos_dir, logo_file)
            if os.path.exists(src_path):
                shutil.copy2(src_path, dst_path)
                data['logos'][company_key] = dst_path
                print(f"  ✓ {company_key}")
            else:
                print(f"  ⚠ Logo not found: {logo_file}")
        
        print("\n✓ Repository content extraction completed successfully")
        return data
        
    except Exception as e:
        print(f"⚠ Error reading repository files: {e}")
        import traceback
        traceback.print_exc()
        return None

# Data
contact_info = {
    "name": "Daniel Hansson",
    "email": "mailto@danielhansson.nu",
    "linkedin": "linkedin.com/in/daniel-hansson-7564a490",
    "github": "github.com/enoch85",
    "location": "Genarp, Skåne, Sweden",
    "website": "rekrytera.danielhansson.nu"
}

# Extract content from repository files
repo_data = extract_content_from_repo()

# Build data structures from extracted content or use fallbacks
if repo_data:
    profile_text_swedish = repo_data.get('profile_text_sv', '')
    profile_text_english = repo_data.get('profile_text_en', '')
    profile_header_swedish = repo_data.get('profile_title_sv', 'DATACENTER MANAGER • OPERATIONS ENGINEER')
    profile_header_english = repo_data.get('profile_title_en', 'DATACENTER MANAGER • OPERATIONS ENGINEER')
    company_logos = repo_data.get('logos', {})
    
    # Convert extracted experiences to the format needed by create_cv
    work_experience_swedish = []
    work_experience_english = []
    
    for exp in repo_data.get('experiences', []):
        # Swedish version
        exp_sv = {
            'title': exp.get('title', ''),
            'company': exp.get('company', ''),
            'location': exp.get('location', ''),
            'date': exp.get('period', ''),
            'description': exp.get('description', ''),
            'keywords': ', '.join(exp.get('skills', [])),
            'page_break_before': exp.get('highlight', '') == 'Karriärutveckling'  # Add page break for certain entries
        }
        work_experience_swedish.append(exp_sv)
        
        # English version (translate)
        exp_en = {
            'title': translate_text(exp.get('title', '')),
            'company': exp.get('company', ''),  # Keep company names unchanged
            'location': exp.get('location', ''),
            'date': exp.get('period', '').replace('Nuvarande', 'Present').replace('Maj', 'May'),
            'description': translate_text(exp.get('description', '')),
            'keywords': ', '.join(exp.get('skills', [])),  # Skills are typically in English already
            'page_break_before': exp.get('highlight', '') == 'Karriärutveckling'
        }
        work_experience_english.append(exp_en)
    
    # Convert extracted skills
    it_skills_swedish = {}
    it_skills_english = {}
    
    for skill_cat in repo_data.get('skills', []):
        title_sv = skill_cat.get('title', '')
        title_en = translate_text(title_sv)
        skills_list = ', '.join(skill_cat.get('skills', []))
        it_skills_swedish[title_sv] = skills_list
        it_skills_english[title_en] = skills_list
    
    # Convert extracted projects
    projects_swedish = []
    projects_english = []
    
    for proj in repo_data.get('projects', []):
        projects_swedish.append({
            'name': proj.get('title', ''),
            'desc': proj.get('description', '')
        })
        projects_english.append({
            'name': proj.get('title', ''),  # Keep project names in original
            'desc': translate_text(proj.get('description', ''))
        })
    
    print("✓ Using dynamically extracted content from repository files")
else:
    # Fallback to hardcoded values
    print("⚠ Using fallback hardcoded content")
    profile_text_swedish = (
        "Med passion för automation, stabilitet och kontinuerlig förbättring. "
        "Kombinerar teknisk expertis inom Linux, virtualisering och infrastruktur med beprövad "
        "förmåga att leda team och driva moderniserings- och migreringsprojekt. "
        "Entreprenöriell bakgrund med affärsförståelse och lösningsorienterat mindset."
    )
    profile_text_english = (
        "With a passion for automation, stability, and continuous improvement. Combines "
        "technical expertise in Linux, virtualization, and infrastructure with proven "
        "ability to lead teams and drive modernization and migration projects. "
        "Entrepreneurial background with business understanding and solution-oriented mindset."
    )
    profile_header_swedish = "DATACENTER MANAGER • OPERATIONS ENGINEER"
    profile_header_english = "DATACENTER MANAGER • OPERATIONS ENGINEER"
    company_logos = {}
    
    # Fallback skills and projects  
    it_skills_swedish = {
        "Virtualisering & Hypervisors": "VMware vSphere/ESX, Nutanix AHV, Hyper-V, Proxmox VE, HCI",
        "OS & Servrar": "Linux (Ubuntu/Debian/Alpine), Windows Server, AD, Exchange, Azure, Docker",
        "Backup & Lagring": "Veeam, Ahsay, Proxmox PBS, Disaster Recovery, HA",
        "Nätverk & Säkerhet": "VPN, DNS/BIND, Fortigate, opnSense, pfSense, Unifi",
        "Automation": "Ansible, Bash, PowerShell, VS Code, IaC, Packer, Git"
    }
    
    it_skills_english = {
        "Virtualization & Hypervisors": "VMware vSphere/ESX, Nutanix AHV, Hyper-V, Proxmox VE, HCI",
        "OS & Servers": "Linux (Ubuntu/Debian/Alpine), Windows Server, AD, Exchange, Azure, Docker",
        "Backup & Storage": "Veeam, Ahsay, Proxmox PBS, Disaster Recovery, HA",
        "Network & Security": "VPN, DNS/BIND, Fortigate, opnSense, pfSense, Unifi",
        "Automation": "Ansible, Bash, PowerShell, VS Code, IaC, Packer, Git"
    }
    
    work_experience_swedish = [
    {
        "title": "Datacenter Manager",
        "company": "InfraCom Smart Digital Solutions",
        "location": "Malmö",
        "date": "Nov 2023 – Nuvarande",
        "description": (
            "Ansvarar för ett team i södra Sverige som jobbar med drift och utveckling av datacenter nationellt. "
            "Säkerställer hög tillgänglighet för kritisk infrastruktur genom proaktiv monitoring, incident management och kontinuerliga förbättringar. "
            "Driver moderniseringsprojekt inkl. migrering mellan hypervisors, och automation av rutinuppgifter. "
            "Jobbar under ISO 27001 enligt ITIL-ramverk."
        ),
        "keywords": "Datacenter Operations, VMware vSphere, Nutanix, Proxmox, Team Leadership, Incident Management"
    },
    {
        "title": "Applikationsspecialist",
        "company": "Redpill Linpro",
        "location": "Linköping (Remote)",
        "date": "Feb 2023 – Nov 2023",
        "page_break_before": True,
        "description": (
            "Implementerade och supporterade enterprise open source-lösningar för svenska myndigheter. "
            "Fokus på säkra, skalbara plattformar baserade på öppna standarder med krav på hög tillgänglighet och compliance."
        ),
        "keywords": "Nextcloud Enterprise, Jitsi, Linux, Customer Support"
    },
    {
        "title": "Founder & CTO",
        "company": "T&M Hansson IT AB",
        "location": "Genarp",
        "date": "Sep 2018 – Nov 2023",
        "description": (
            "Grundade och drev IT-konsultbolag specialiserat på open source-infrastruktur och automation. "
            "Maintainer för Nextclouds officiella VM (100+ nedladdningar/dag globalt). "
            "Utvecklade automatiserade deployment-lösningar för flera virtualiseringsplattformar. "
            "Supporterade kunder över hela världen. Sålde bolaget till Kafit AB 2021."
        ),
        "keywords": "Entrepreneurship, IaC, Bash Automation, Open Source Maintenance"
    },
    {
        "title": "Operations / DevOps Engineer",
        "company": "SenseNode",
        "location": "Lund",
        "date": "Feb 2022 – Feb 2023",
        "description": (
            "Tillverkade och installerade IoT-enheter med Ansible för energimätning hos stora industrier. "
            "Ansvarade för att organisera RAM-avtal med telefonioperatörer och säkerställa enheter alltid var uppkopplade. "
            "Gav support och förvaltade infrastruktur för säker dataöverföring."
        ),
        "keywords": "IoT, Operations, Customer Support, Vendor Management"
    },
    {
        "title": "IT-tekniker → Systemadministratör",
        "company": "Vessinge IT",
        "location": "Malmö",
        "date": "Jul 2017 – Jan 2022",
        "description": (
            "Började som IT-tekniker och avancerade till systemadministratör med huvudansvar för alla servrar och mjukvaror. "
            "Blev senare ansvarig för Second-Line support. "
            "Förvaltade Windows/Linux hybrid-miljö med VMware, Hyper-V, Windows Server 2008-2019, Ubuntu, Docker och Portainer. "
            "Daglig kontakt med leverantörer och åtgärdade problem på högre teknisk nivå."
        ),
        "keywords": "System Admin, VMware/Hyper-V, Windows/Linux Hybrid, SQL Server"
    },
    {
        "title": "Account Manager",
        "company": "Örestad Linux AB",
        "location": "Malmö",
        "date": "Nov 2016 – Maj 2017",
        "description": (
            "Teknisk säljroll med fokus på Linux-baserade hosting-lösningar. "
            "Gjorde Örestad Linux till Nextcloud partner och implementerade VPS-plattform för Nextcloud-hosting."
        ),
        "keywords": "Technical Pre-Sales, Linux Solutions, Nextcloud",
        "page_break_before": True
    },
    {
        "title": "Tidigare entreprenörskarriär och försäljning",
        "company": "",
        "location": "Västerås, Malmö",
        "date": "2008 – 2016",
        "description": (
            "HSN Konsult (2008-2011): Grundare, B2B-försäljning. Sålde företaget 2011.\n"
            "Powersales Sweden AB (2011-2014): CTO. Ansvarade för IT.\n"
            "Trygg-Hansa (2014-2016): Försäkringsrådgivare."
        ),
        "keywords": "Business Development, CTO, B2B Sales"
    }
]

work_experience_english = [
    {
        "title": "Datacenter Manager",
        "company": "InfraCom Smart Digital Solutions",
        "location": "Malmö",
        "date": "Nov 2023 – Present",
        "description": (
            "Responsible for a team in southern Sweden working with operations and development of datacenters nationally. "
            "Ensures high availability for critical infrastructure through proactive monitoring, incident management, and continuous improvements. "
            "Drives modernization projects including migration between hypervisors, and automation of routine tasks. "
            "Working under ISO 27001 according to ITIL framework."
        ),
        "keywords": "Datacenter Operations, VMware vSphere, Nutanix, Proxmox, Team Leadership, Incident Management"
    },
    {
        "title": "Application Specialist",
        "company": "Redpill Linpro",
        "location": "Linköping (Remote)",
        "date": "Feb 2023 – Nov 2023",
        "page_break_before": True,
        "description": (
            "Implemented and supported enterprise open source solutions for Swedish government agencies. "
            "Focus on secure, scalable platforms based on open standards with requirements for high availability and compliance."
        ),
        "keywords": "Nextcloud Enterprise, Jitsi, Linux, Customer Support"
    },
    {
        "title": "Founder & CTO",
        "company": "T&M Hansson IT AB",
        "location": "Genarp",
        "date": "Sep 2018 – Nov 2023",
        "description": (
            "Founded and ran an IT consulting company specialized in open source infrastructure and automation. "
            "Maintainer for Nextcloud's official VM (100+ downloads/day globally). "
            "Developed automated deployment solutions for several virtualization platforms. "
            "Supported customers worldwide. Sold the company to Kafit AB 2021."
        ),
        "keywords": "Entrepreneurship, IaC, Bash Automation, Open Source Maintenance"
    },
    {
        "title": "Operations / DevOps Engineer",
        "company": "SenseNode",
        "location": "Lund",
        "date": "Feb 2022 – Feb 2023",
        "description": (
            "Manufactured and installed IoT devices with Ansible for energy metering at large industries. "
            "Responsible for organizing framework agreements with telephony operators and ensuring devices were always connected. "
            "Provided support and managed infrastructure for secure data transfer."
        ),
        "keywords": "IoT, Operations, Customer Support, Vendor Management"
    },
    {
        "title": "IT Technician → System Administrator",
        "company": "Vessinge IT",
        "location": "Malmö",
        "date": "Jul 2017 – Jan 2022",
        "description": (
            "Started as IT technician and advanced to system administrator with main responsibility for all servers and software. "
            "Later became responsible for Second-Line support. "
            "Managed Windows/Linux hybrid environment with VMware, Hyper-V, Windows Server 2008-2019, Ubuntu, Docker, and Portainer. "
            "Daily contact with vendors and resolved issues at a higher technical level."
        ),
        "keywords": "System Admin, VMware/Hyper-V, Windows/Linux Hybrid, SQL Server"
    },
    {
        "title": "Account Manager",
        "company": "Örestad Linux AB",
        "location": "Malmö",
        "date": "Nov 2016 – May 2017",
        "description": (
            "Technical sales role focusing on Linux-based hosting solutions. "
            "Made Örestad Linux a Nextcloud partner and implemented VPS platform for Nextcloud hosting."
        ),
        "keywords": "Technical Pre-Sales, Linux Solutions, Nextcloud",
        "page_break_before": True
    },
    {
        "title": "Previous Sales Career",
        "company": "",
        "location": "Västerås, Malmö",
        "date": "2008 – 2016",
        "description": (
            "HSN Konsult (2008-2011): Founder, B2B sales. Sold company in 2011.\n"
            "Powersales Sweden AB (2011-2014): CTO. Responsible for IT.\n"
            "Trygg-Hansa (2014-2016): Insurance Advisor."
        ),
        "keywords": "Business Development, CTO, B2B Sales"
    }
]

# Hardcoded data for competencies, education, certificates, and languages (same for both success/fallback paths)
key_competencies_swedish = [
    "Datacenterdrift & Management",
    "Server- & Infrastrukturadministration",
    "Virtualisering & Hypervisors (VMware, Nutanix, Proxmox, Hyper-V)",
    "Open Source & Automation",
    "Entreprenörskap & Affärsutveckling"
]

key_competencies_english = [
    "Datacenter Operations & Management",
    "Server & Infrastructure Administration",
    "Virtualization & Hypervisors (VMware, Nutanix, Proxmox, Hyper-V)",
    "Open Source & Automation",
    "Entrepreneurship & Business Development"
]

education_swedish = [
    {
        "degree": "IT & Sales (300 YH-poäng)",
        "school": "EC Utbildning, Malmö",
        "date": "2013 – 2014"
    },
    {
        "degree": "Gymnasieexamen",
        "school": "Västerås",
        "date": ""
    }
]

education_english = [
    {
        "degree": "IT & Sales (300 YH-points)",
        "school": "EC Utbildning, Malmö",
        "date": "2013 – 2014"
    },
    {
        "degree": "High School Diploma",
        "school": "Västerås",
        "date": ""
    }
]

certificates_swedish = [
    "Nutanix ECA (2024)",
    "Certifierad Försäkringsförmedlare (2014)",
    "Microsoft Technical Advisor (2016)"
]

certificates_english = [
    "Nutanix ECA (2024)",
    "Certified Insurance Intermediary (2014)",
    "Microsoft Technical Advisor (2016)"
]

projects_swedish = [
    {
        "name": "Nextcloud VM",
        "desc": "Officiell Nextcloud-appliance med interaktiva skript för säker hosting. Används globalt."
    },
    {
        "name": "OVMS Home Assistant",
        "desc": "MQTT-integration för Open Vehicle Monitoring System."
    },
    {
        "name": "Global Energy Spotprices",
        "desc": "Global elprisprognos-integration."
    },
    {
        "name": "EffektGuard",
        "desc": "Intelligent värmepumpsoptimering med ML."
    }
]

projects_english = [
    {
        "name": "Nextcloud VM",
        "desc": "Official Nextcloud appliance with interactive scripts for secure hosting. Used globally."
    },
    {
        "name": "OVMS Home Assistant",
        "desc": "MQTT integration for Open Vehicle Monitoring System."
    },
    {
        "name": "Global Energy Spotprices",
        "desc": "Global electricity price forecast integration."
    },
    {
        "name": "EffektGuard",
        "desc": "Intelligent heat pump optimization using ML."
    }
]

languages_swedish = [
    "Svenska (Modersmål)",
    "Engelska (Flytande)",
    "Portugisiska (Nybörjare)"
]

languages_english = [
    "Swedish (Native)",
    "English (Fluent)",
    "Portuguese (Beginner)"
]

def set_table_borders_to_none(table):
    """
    Removes borders from a table by setting all border elements to nil.
    Also attempts to clear cell borders.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    # 1. Remove tblStyle to prevent style inheritance
    tblStyle = tblPr.first_child_found_in("w:tblStyle")
    if tblStyle is not None:
        tblPr.remove(tblStyle)

    # 2. Remove tblLook to prevent conditional formatting
    tblLook = tblPr.first_child_found_in("w:tblLook")
    if tblLook is not None:
        tblPr.remove(tblLook)
    
    # 3. Create and set Table Borders (tblBorders)
    # Remove existing if any
    existing_borders = tblPr.first_child_found_in("w:tblBorders")
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'FFFFFF')
        tblBorders.append(border)
    
    # Insert tblBorders in correct position
    # Order: tblW -> tblBorders -> tblLayout
    # We will append it, but if tblLayout exists, insert before it.
    
    tblLayout = tblPr.first_child_found_in("w:tblLayout")
    if tblLayout is not None:
        tblLayout.addprevious(tblBorders)
    else:
        tblPr.append(tblBorders)

    # 4. Cell Borders (tcBorders) - Iterate through all cells
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is not None:
                tcPr.remove(tcBorders)
            
            # Explicitly set nil borders for cell
            # NOTE: tcBorders does NOT support insideH/insideV, only top/left/bottom/right/tl2br/tr2bl
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'FFFFFF')
                tcBorders.append(border)

def create_cv(template_path, output_path, image_path, language='en'):
    """
    Create a CV document in the specified language.
    
    Args:
        template_path: Path to the Word template
        output_path: Where to save the generated CV
        image_path: Path to profile image
        language: 'sv' for Swedish or 'en' for English
    """
    doc = Document(template_path)
    
    # Select data based on language
    if language == 'sv':
        profile_header = profile_header_swedish
        profile_text = profile_text_swedish
        key_competencies = key_competencies_swedish
        it_skills = it_skills_swedish
        work_experience = work_experience_swedish
        education = education_swedish
        certificates = certificates_swedish
        projects = projects_swedish
        languages_list = languages_swedish
        work_exp_header = "ARBETSLIVSERFARENHET"
        competencies_header = "NYCKELKOMPETENSER"
        education_header = "UTBILDNING"
        certificates_header = "Certifieringar:"
        skills_header = "IT-KOMPETENS"
        languages_header = "SPRÅK"
        projects_header = "PROJEKT & OPEN SOURCE"
    else:  # English
        profile_header = profile_header_english
        profile_text = profile_text_english
        key_competencies = key_competencies_english
        it_skills = it_skills_english
        work_experience = work_experience_english
        education = education_english
        certificates = certificates_english
        projects = projects_english
        languages_list = languages_english
        work_exp_header = "WORK EXPERIENCE"
        competencies_header = "KEY COMPETENCIES"
        education_header = "EDUCATION"
        certificates_header = "Certificates:"
        skills_header = "IT SKILLS"
        languages_header = "LANGUAGES"
        projects_header = "PROJECTS & OPEN SOURCE"
    
    # --- 1. Remove the existing table entirely ---
    if len(doc.tables) > 0:
        old_table = doc.tables[0]
        old_table._element.getparent().remove(old_table._element)
    
    # --- 2. Clear ALL existing paragraphs after the table ---
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    
    # --- 3. Create header using a paragraph with floating image ---
    
    # Name (large, bold) - First paragraph with image
    p = doc.add_paragraph()
    
    # Add image first with text wrapping (Square/Tight wrapping, anchored to right)
    if os.path.exists(image_path):
        run = p.add_run()
        inline_shape = run.add_picture(image_path, width=Inches(1.9))
        
        # Convert inline image to floating (anchored) image
        # This allows text to wrap around it
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        
        # Get the inline shape element
        inline = inline_shape._inline
        
        # Create anchor element for floating image - positioned at top right with margin
        anchor_xml = (
            '<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
            'distT="0" distB="0" distL="114300" distR="114300" simplePos="0" relativeHeight="251658240" '
            'behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">'
            '<wp:simplePos x="0" y="0"/>'
            '<wp:positionH relativeFrom="page">'
            '<wp:posOffset>4900000</wp:posOffset>'
            '</wp:positionH>'
            '<wp:positionV relativeFrom="page">'
            '<wp:posOffset>900000</wp:posOffset>'
            '</wp:positionV>'
            '<wp:extent cx="{cx}" cy="{cy}"/>'
            '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
            '<wp:wrapSquare wrapText="left"/>'
            '<wp:docPr id="1" name="Picture 1"/>'
            '<wp:cNvGraphicFramePr/>'
            '{graphic}'
            '</wp:anchor>'
        ).format(
            cx=inline.extent.cx,
            cy=inline.extent.cy,
            graphic=inline.graphic.xml.decode() if isinstance(inline.graphic.xml, bytes) else str(inline.graphic.xml)
        )
        
        anchor = parse_xml(anchor_xml)
        inline.getparent().replace(inline, anchor)
    
    # Now add the text content
    runner = p.add_run(contact_info["name"])
    runner.bold = True
    runner.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(0)
    p.add_run('\n')
    
    # Contact info - all in one paragraph with line breaks
    p.add_run('Email: ')
    add_hyperlink(p, 'mailto:mailto@danielhansson.nu', 'mailto@danielhansson.nu')
    p.add_run('\n')
    
    p.add_run('Telephone: ')
    add_hyperlink(p, 'tel:+46734045555', '+46 734 04 55 55')
    p.add_run('\n')
    
    p.add_run('LinkedIn: ')
    add_hyperlink(p, 'https://linkedin.com/in/daniel-hansson-7564a490', 'linkedin.com/in/daniel-hansson-7564a490')
    p.add_run('\n')
    
    p.add_run('GitHub: ')
    add_hyperlink(p, 'https://github.com/enoch85', 'github.com/enoch85')
    p.add_run('\n')
    
    p.add_run('Website: ')
    add_hyperlink(p, 'https://rekrytera.danielhansson.nu', 'rekrytera.danielhansson.nu')
    p.add_run('\n')
    
    p.add_run('Location: ')
    p.add_run(contact_info['location'])
    p.paragraph_format.space_after = Pt(12)
        
    # --- 3. Rebuild Content (Compact) ---
    
    def add_header(text):
        p = doc.add_paragraph(text)
        p.style = 'Heading 2'
        p.runs[0].bold = True
        p.runs[0].font.all_caps = True
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        return p

    # Profile
    add_header(profile_header)
    p = doc.add_paragraph(profile_text)
    p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph("").paragraph_format.space_after = Pt(6)

    # Key Competencies
    add_header(competencies_header)
    for comp in key_competencies:
        p = doc.add_paragraph(comp)
        try:
            p.style = 'List Bullet'
        except KeyError:
            p.text = f"• {comp}"
        p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph("").paragraph_format.space_after = Pt(6)

    # Work Experience
    add_header(work_exp_header)
    
    # Download logos on the fly for each job
    for job in work_experience:
        # Check if we have a logo for this company - download if needed
        logo_path = None
        # Map "Previous Sales Career" to HSN Konsult logo
        if 'Previous Sales Career' in job['title']:
            company_key = 'HSN Konsult'
            if company_key not in company_logos:
                # Download HSN Konsult logo
                try:
                    response_web = requests.get("https://rekrytera.danielhansson.nu", timeout=5)
                    soup_web = BeautifulSoup(response_web.text, 'html.parser')
                    imgs = soup_web.find_all('img', alt=True)
                    for img in imgs:
                        if company_key.lower() in img.get('alt', '').lower():
                            src = img.get('src')
                            if src:
                                ext = src.split('.')[-1] if '.' in src else 'png'
                                logo_file = company_key.lower().replace(' ', '-') + f'-logo.{ext}'
                                downloaded = download_logo(src, logo_file)
                                if downloaded:
                                    company_logos[company_key] = downloaded
                                    print(f"  ✓ Downloaded {company_key} logo")
                                break
                except:
                    pass
            
            if company_key in company_logos:
                logo_path = company_logos[company_key]
        else:
            # Regular company matching
            for company_key in ['InfraCom', 'Redpill Linpro', 'T&M Hansson IT', 'SenseNode', 'Vessinge IT', 'Örestad Linux']:
                # For Örestad, also check the 'Orestad Linux' key used in website alt text
                logo_key = 'Orestad Linux' if company_key == 'Örestad Linux' else company_key
                
                if company_key.lower() in job['company'].lower():
                    # Check if already in company_logos dict (using website's alt text key)
                    if logo_key in company_logos:
                        logo_path = company_logos[logo_key]
                    elif company_key not in company_logos:
                        # Try to load from file or download
                        logo_filename = 'orestad-logo.png' if company_key == 'Örestad Linux' else company_key.lower().replace(' ', '-') + '-logo.png'
                        logo_filepath = os.path.join(os.path.dirname(__file__), '..', 'logos', logo_filename)
                        if os.path.exists(logo_filepath):
                            company_logos[company_key] = logo_filepath
                            logo_path = logo_filepath
                            print(f"  ✓ Using existing {company_key} logo")
                        else:
                            # Try to download from website
                            try:
                                response_web = requests.get("https://rekrytera.danielhansson.nu", timeout=5)
                                soup_web = BeautifulSoup(response_web.text, 'html.parser')
                                imgs = soup_web.find_all('img', alt=True)
                                for img in imgs:
                                    if company_key.lower() in img.get('alt', '').lower():
                                        src = img.get('src')
                                        if src:
                                            ext = src.split('.')[-1] if '.' in src else 'png'
                                            logo_file = company_key.lower().replace(' ', '-') + f'-logo.{ext}'
                                            downloaded = download_logo(src, logo_file)
                                            if downloaded:
                                                company_logos[company_key] = downloaded
                                                logo_path = downloaded
                                                print(f"  ✓ Downloaded {company_key} logo")
                                            break
                            except:
                                pass
                    break
        
        # Add page break if specified
        if job.get('page_break_before', False):
            doc.add_page_break()
        
        # Create paragraph with logo
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        
        # Add logo if available (bigger size)
        if logo_path and os.path.exists(logo_path):
            try:
                run = p.add_run()
                run.add_picture(logo_path, height=Inches(0.35))  # Bigger logo
                p.add_run("  ")  # Space after logo
            except Exception as e:
                print(f"  ⚠ Could not add logo for {job['company']}: {e}")
        
        # Job title on same line as logo
        r1 = p.add_run(f"{job['title']}")
        r1.bold = True
        
        # Company, location, date on new line (skip if empty)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        info_parts = []
        if job['company']:
            info_parts.append(job['company'])
        if job['location']:
            info_parts.append(job['location'])
        if job['date']:
            info_parts.append(job['date'])
        if info_parts:
            p.add_run(' | '.join(info_parts))
        else:
            # If no company/location, just show date
            p.add_run(job['date'])
        
        # Description
        p = doc.add_paragraph(job['description'])
        p.paragraph_format.space_after = Pt(0)
        
        # Keywords
        p = doc.add_paragraph()
        r = p.add_run(f"Keywords: {job['keywords']}")
        r.italic = True
        r.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(0)
    
    doc.add_paragraph("").paragraph_format.space_after = Pt(6)

    # Education
    add_header(education_header)
    for edu in education:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"{edu['degree']}")
        r.bold = True
        # Add school and date, handling empty values
        if edu['school'] and edu['date']:
            p.add_run(f", {edu['school']} ({edu['date']})")
        elif edu['school']:
            p.add_run(f", {edu['school']}")
        elif edu['date']:
            p.add_run(f" ({edu['date']})")
    doc.add_paragraph("").paragraph_format.space_after = Pt(6)
    
    # Certificates (Bullet Points)
    p = doc.add_paragraph(certificates_header)
    p.runs[0].bold = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    
    for cert in certificates:
        p = doc.add_paragraph(cert)
        try:
            p.style = 'List Bullet'
        except KeyError:
            p.text = f"• {cert}"
        p.paragraph_format.space_after = Pt(0)
    
    doc.add_paragraph("").paragraph_format.space_after = Pt(6)

    # IT Skills
    add_header(skills_header)
    for cat, skills in it_skills.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"{cat}: ")
        r.bold = True
        p.add_run(skills)

    # Languages (Bullet Points)
    add_header(languages_header)
    for l in languages_list:
        p = doc.add_paragraph(l)
        try:
            p.style = 'List Bullet'
        except KeyError:
            p.text = f"• {l}"
        p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph("").paragraph_format.space_after = Pt(6)

    # Projects
    add_header(projects_header)
    for proj in projects:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"{proj['name']}")
        r.bold = True
        p.add_run(f": {proj['desc']}")
    doc.add_paragraph("").paragraph_format.space_after = Pt(6)

    # Add page numbering to header
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add page number field
    run = header_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)
    
    # Add " (" text
    header_para.add_run(' (')
    
    # Add total pages field
    run2 = header_para.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    run2._r.append(fldChar3)
    
    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = 'NUMPAGES'
    run2._r.append(instrText2)
    
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run2._r.append(fldChar4)
    
    # Add ")" text
    header_para.add_run(')')

    doc.save(output_path)
    print(f"✓ CV saved to {output_path}")

if __name__ == "__main__":
    import argparse
    
    # Get script directory to build relative paths
    script_dir = os.path.dirname(os.path.abspath(__file__))
    repo_root = os.path.dirname(script_dir)
    
    parser = argparse.ArgumentParser(description='Generate CV in Swedish and/or English')
    parser.add_argument('--template', default=os.path.join(script_dir, 'mall-kronologiskt-cv-251020-variant.docx'),
                       help='Path to Word template')
    parser.add_argument('--image', default=os.path.join(script_dir, 'profile_pic.jpg'),
                       help='Path to profile image')
    parser.add_argument('--output-dir', default=script_dir,
                       help='Directory to save CVs')
    parser.add_argument('--lang', choices=['sv', 'en', 'both'], default='both',
                       help='Language(s) to generate: sv, en, or both')
    
    args = parser.parse_args()
    
    print("\n" + "="*60)
    print("  CV GENERATOR - Multi-language Edition")
    print("="*60)
    
    if args.lang in ['sv', 'both']:
        print("\n🇸🇪 Generating Swedish CV...")
        output_sv = os.path.join(args.output_dir, 'Daniel_Hansson_CV_2025_SV.docx')
        create_cv(args.template, output_sv, args.image, language='sv')
    
    if args.lang in ['en', 'both']:
        print("\n🇬🇧 Generating English CV...")
        output_en = os.path.join(args.output_dir, 'Daniel_Hansson_CV_2025_ENG.docx')
        create_cv(args.template, output_en, args.image, language='en')
    
    print("\n" + "="*60)
    print("  ✓ CV Generation Complete!")
    print("="*60)
